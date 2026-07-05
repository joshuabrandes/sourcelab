import logging
from pathlib import Path

from docx import Document

from sidecar.extractors.markitdown_adapter import convert_to_markdown
from sidecar.extractors.utils import markdown_to_elements, rows_to_markdown_table, utc_timestamp
from sidecar.models import ContentType, DocumentElement, DocumentMetadata, ElementType, ExtractedDocument

_logger = logging.getLogger(__name__)


def extract_docx_document(source_id: str, file_path: str) -> ExtractedDocument:
    path = Path(file_path).expanduser().resolve()
    if not path.exists() or not path.is_file():
        raise FileNotFoundError(f"File not found at {path}")
    if path.suffix.lower() != ".docx":
        raise ValueError(f"Unsupported file type: {path.suffix}")

    document = Document(str(path))
    try:
        markdown, converted_title = convert_to_markdown(path)
        elements = markdown_to_elements(markdown)
    except Exception as exc:
        _logger.warning("MarkItDown extraction failed for %s; using python-docx fallback: %s", path, exc)
        converted_title = None
        elements = _extract_docx_fallback(document)

    properties = document.core_properties
    title = converted_title or properties.title or path.stem
    metadata = DocumentMetadata(
        author=properties.author or None,
        createdAt=properties.created.isoformat() if properties.created else None,
        extractedAt=utc_timestamp(),
    )
    return ExtractedDocument(
        sourceId=source_id,
        title=title,
        language=None,
        contentType=ContentType.docx,
        metadata=metadata,
        elements=elements,
    )


def _extract_docx_fallback(document: Document) -> list[DocumentElement]:
    elements: list[DocumentElement] = []
    paragraphs = {paragraph._element: paragraph for paragraph in document.paragraphs}
    tables = {table._element: table for table in document.tables}

    for block in document.element.body:
        tag = block.tag.rsplit("}", maxsplit=1)[-1]
        if tag == "p":
            paragraph = paragraphs.get(block)
            if paragraph is None or not paragraph.text.strip():
                continue
            style_name = paragraph.style.name if paragraph.style else ""
            level = _heading_level(style_name)
            elements.append(
                DocumentElement(
                    type=ElementType.heading if level else ElementType.paragraph,
                    content=paragraph.text.strip(),
                    level=level,
                    position=len(elements),
                )
            )
        elif tag == "tbl":
            table = tables.get(block)
            if table is None:
                continue
            elements.append(
                DocumentElement(
                    type=ElementType.table,
                    content=_table_to_markdown(table),
                    position=len(elements),
                )
            )

    return elements or [DocumentElement(type=ElementType.paragraph, content="", position=0)]


def _heading_level(style_name: str) -> int | None:
    if not style_name.lower().startswith("heading"):
        return None
    try:
        return max(1, min(int(style_name.rsplit(maxsplit=1)[-1]), 6))
    except ValueError:
        return 1


def _table_to_markdown(table) -> str:
    return rows_to_markdown_table([cell.text for cell in row.cells] for row in table.rows)
