from docx import Document

from sidecar.extractors.docx import extract_docx_document
from sidecar.models import ContentType, ElementType


def create_docx(path):
    document = Document()
    document.core_properties.author = "Ada Lovelace"
    document.core_properties.title = "Architecture Notes"
    document.add_heading("Overview", level=1)
    document.add_paragraph("Fallback paragraph")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "A"
    table.cell(0, 1).text = "B"
    table.cell(1, 0).text = "1"
    table.cell(1, 1).text = "2"
    document.save(path)


def test_docx_uses_markitdown_and_preserves_metadata(tmp_path, mocker):
    path = tmp_path / "notes.docx"
    create_docx(path)
    convert = mocker.patch(
        "sidecar.extractors.docx.convert_to_markdown",
        return_value=("# Converted\n\nMain content.", None),
    )

    extracted = extract_docx_document("source-1", str(path))

    convert.assert_called_once_with(path.resolve())
    assert extracted.contentType == ContentType.docx
    assert extracted.title == "Architecture Notes"
    assert extracted.metadata.author == "Ada Lovelace"
    assert extracted.metadata.pageCount is None
    assert [element.type for element in extracted.elements] == [
        ElementType.heading,
        ElementType.paragraph,
    ]


def test_docx_falls_back_to_structured_python_docx(tmp_path, mocker):
    path = tmp_path / "fallback.docx"
    create_docx(path)
    mocker.patch("sidecar.extractors.docx.convert_to_markdown", side_effect=RuntimeError("failed"))

    extracted = extract_docx_document("source-2", str(path))

    assert [element.type for element in extracted.elements] == [
        ElementType.heading,
        ElementType.paragraph,
        ElementType.table,
    ]
    assert extracted.elements[0].content == "Overview"
    assert "| A | B |" in extracted.elements[2].content


def test_docx_markitdown_integration(tmp_path):
    path = tmp_path / "integration.docx"
    create_docx(path)

    extracted = extract_docx_document("source-3", str(path))

    assert any(element.type == ElementType.heading for element in extracted.elements)
    assert any(element.type == ElementType.paragraph for element in extracted.elements)
    assert any(element.type == ElementType.table for element in extracted.elements)
